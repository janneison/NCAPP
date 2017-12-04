# -*- coding: utf-8 -*-

from __future__ import unicode_literals

from django.shortcuts import render
""" 
# Create your views here.
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH , WD_COLOR_INDEX
from rest_framework.request import Request

from levantamiento.models import (Proyecto, Convenio, 
Circuito, Apoyo, Version, TipoDato, Capitulo, 
Levantamiento, Estado, DetalleLevantamiento,SoporteNC,
Norma, Defecto, Actividad)

@login_required
def generar_informe(request):
    document = Document()
    paragraph = document.add_paragraph('Informe de Calidad.')

    document.add_heading('Calidad de las obras')

    document.add_paragraph('A continuación, se presentan los hallazgos representativos de calidad hallados en las inspecciones de calidad:')


    table = document.add_table(rows=2, cols=2)

    document.save('informe/test.docx')
""" 